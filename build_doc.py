from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
BLACK     = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
GOLD      = RGBColor(0xC9, 0xA8, 0x00) # deep gold
DARK_GOLD = RGBColor(0xA0, 0x82, 0x00)
CREAM     = RGBColor(0xFA, 0xF7, 0xF0)
GREY_MID  = RGBColor(0x55, 0x55, 0x55)
GREY_LIGHT= RGBColor(0xDD, 0xDD, 0xDD)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph background shading ──────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# ── Helper: set cell background ───────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add a thick top border to a paragraph ────────────────────────────
def top_border(para, hex_color='C9A800', width_pt=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    str(width_pt))
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), hex_color)
    pBdr.append(top)
    pPr.append(pBdr)

# ── Helper: remove spacing before/after ──────────────────────────────────────
def tight(para, before=0, after=60):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

# ── Helper: horizontal rule ───────────────────────────────────────────────────
def hr(doc, color='C9A800'):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════

# Big title
p = doc.add_paragraph()
shade_paragraph(p, '1A1A2E')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(p, before=0, after=4)
run = p.add_run('ONEDESIGN.PL')
run.bold      = True
run.font.size = Pt(36)
run.font.color.rgb = GOLD
run.font.name = 'Calibri'

p2 = doc.add_paragraph()
shade_paragraph(p2, '1A1A2E')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(p2, before=0, after=4)
r2 = p2.add_run('INSTAGRAM DOMINATION MASTERPLAN')
r2.bold = True
r2.font.size = Pt(16)
r2.font.color.rgb = WHITE
r2.font.name = 'Calibri'

p3 = doc.add_paragraph()
shade_paragraph(p3, '1A1A2E')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(p3, before=0, after=16)
r3 = p3.add_run('Written from the perspective of your CEO  ·  No fluff. No links. Everything you need is here.')
r3.italic = True
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0xBB, 0xAA, 0x80)
r3.font.name = 'Calibri'

doc.add_paragraph()  # spacer

# ── Reusable heading functions ────────────────────────────────────────────────
def part_heading(doc, text):
    """Dark band full-width part heading."""
    p = doc.add_paragraph()
    shade_paragraph(p, '1A1A2E')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tight(p, before=18, after=4)
    r = p.add_run(f'  {text}')
    r.bold      = True
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD
    r.font.name = 'Calibri'

def section_heading(doc, text):
    """Gold-bordered section heading."""
    p = doc.add_paragraph()
    top_border(p, 'C9A800', 12)
    tight(p, before=16, after=2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLACK
    r.font.name = 'Calibri'

def sub_heading(doc, text):
    p = doc.add_paragraph()
    tight(p, before=10, after=2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_GOLD
    r.font.name = 'Calibri'

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    tight(p, before=1, after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    r.font.name = 'Calibri'
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    tight(p, before=1, after=3)
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    r.font.name = 'Calibri'

def callout(doc, label, text, bg='FAF7F0', border='C9A800'):
    """Highlighted callout box using a 1-cell table."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg)
    cell.width = Inches(5.5)
    # border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '12')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), border)
        tcBdr.append(el)
    tcPr.append(tcBdr)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rl = p.add_run(f'{label}  ')
    rl.bold = True
    rl.font.color.rgb = DARK_GOLD
    rl.font.size = Pt(10)
    rl.font.name = 'Calibri'
    rb = p.add_run(text)
    rb.font.size = Pt(10)
    rb.font.color.rgb = BLACK
    rb.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, header_bg='1A1A2E', alt_bg='FAF7F0'):
    col_count = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = GOLD
        r.font.size = Pt(9.5)
        r.font.name = 'Calibri'
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
    # data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        bg  = alt_bg if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.color.rgb = BLACK
            r.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# FOREWORD
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'FOREWORD')
body(doc, 'Most interior design studios on Instagram fail for the same reason: they treat Instagram like a digital photo album. They post a beautiful room, write "Nowy projekt ✨" and wonder why nobody follows them.')
body(doc, 'Here is the truth: nobody follows a gallery. People follow a person, a perspective, a promise.')
callout(doc, '→ KEY MINDSET SHIFT:', 'From this day forward, you are not just an interior design studio. You are a media company that also does interior design.')

# ══════════════════════════════════════════════════════════════════════════════
# PART 1
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 1 — THE FOUNDATION')

section_heading(doc, '1.1  The Brand Identity Decision')
body(doc, 'Before you post a single piece of content, you must answer this with total clarity:')
callout(doc, '?', '"Why should someone follow @onedesignpl instead of the 10,000 other interior design accounts in Poland?"')

sub_heading(doc, 'OPTION A — The Transformation Studio')
bullet(doc, 'Core promise: We turn the worst apartments in Warsaw into the best ones.')
bullet(doc, 'Content angle: Drama. Before and afters. "You won\'t believe this is the same room."')
bullet(doc, 'Who it attracts: People with problem spaces, renovation dreamers, young couples.')
bullet(doc, 'Why it works: Before/after is the most shared, most saved, most viral format in interior design.')

sub_heading(doc, 'OPTION B — The Smart Design Studio')
bullet(doc, 'Core promise: We help you make the right decisions so you don\'t waste money.')
bullet(doc, 'Content angle: Education. Mistakes to avoid. Industry secrets. How-tos.')
bullet(doc, 'Who it attracts: People planning a renovation who are anxious about getting it wrong.')
bullet(doc, 'Why it works: Educational content drives the highest saves — a top algorithmic signal.')

sub_heading(doc, 'OPTION C — The Aesthetic Studio  ✗  NOT RECOMMENDED')
body(doc, 'Pure inspiration content. This is what everyone else is doing. You need 50K+ followers before pure aesthetic content builds momentum. Without a unique angle, you are invisible.')

callout(doc, '★ CEO RECOMMENDATION:', 'Start with Option A (Transformation Studio), layer in Option B (Smart Design) by Month 2. Discovery through drama. Conversion through intelligence.')

section_heading(doc, '1.2  Your Three Target People')
body(doc, 'Stop thinking in demographics. Think in specific people.')

sub_heading(doc, 'KASIA, 31 — Your Best Potential Client')
body(doc, 'Just bought a 48m² apartment in Mokotów with her partner. Budget: 80,000 PLN. Excited but terrified. Doesn\'t know what style she wants. Goes to Instagram for inspiration and ends up more confused. Every post you make should make Kasia feel: "These people understand exactly what I\'m going through."')

sub_heading(doc, 'MICHAŁ, 38 — Your Long-Term Investor')
body(doc, 'Owns a 90m² apartment and will renovate in 2 years. Not ready to buy but actively researching. He saves posts. Reads captions carefully. He will follow you today and DM you in 14 months. Build your educational content for him. When he is ready, you will be the only studio he considers.')

sub_heading(doc, 'ANNA, 26 — Your Distribution Engine')
body(doc, 'Will never hire you. Lives outside Warsaw. But she saves your posts, shares your Reels, sends your content to everyone she knows. Her engagement signals the algorithm that your content is valuable, which pushes it to Kasia and Michał. Never dismiss Anna.')

section_heading(doc, '1.3  Brand Voice Rules')
bullet(doc, 'Direct — say what you mean, no corporate speak')
bullet(doc, 'Confident but not arrogant — demonstrate expertise, don\'t brag about it')
bullet(doc, 'Human — have opinions, disagree with trends, be a real point of view')
bullet(doc, 'Helpful first — give before you ask, every post teaches something before asking for anything')
body(doc, '')
bullet(doc, '✗  Never write: "Nowy projekt! Co myślicie? ✨"')
bullet(doc, '✓  Instead write like you\'re texting a smart friend with a strong opinion')

# ══════════════════════════════════════════════════════════════════════════════
# PART 2
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 2 — THE CONTENT MACHINE')

section_heading(doc, '2.1  How the Algorithm Actually Works')
body(doc, 'Instagram\'s algorithm answers one question: "Will this content keep people on Instagram longer?" Here are the five signals it uses, ranked by importance:')

add_table(doc,
    ['Rank', 'Signal', 'What it means', 'How to trigger it'],
    [
        ['1 ★', 'DM Shares', 'Strongest signal for reaching non-followers', 'Create content people send to a specific friend'],
        ['2', 'Saves', 'Content valuable enough to keep', 'Educational posts: mistakes, checklists, guides'],
        ['3', 'Watch Time', '% of Reel watched — quality over views', 'Hook in first 1.5s; cut every unnecessary second'],
        ['4', 'Comments', 'Signals discussion and engagement', 'Ask specific easy-to-answer questions'],
        ['5', 'Likes', 'Weakest signal — do not optimise for this', '—'],
    ]
)

section_heading(doc, '2.2  The Three Content Pillars')

add_table(doc,
    ['Pillar', 'Goal', 'Best Formats', 'Frequency'],
    [
        ['DISCOVERY', 'Reach people who don\'t know you exist', 'Reels, bold carousels', '3× per week'],
        ['TRUST', 'Convert followers into warm leads', 'Carousels, long Reels, Stories', '2× per week'],
        ['CONVERSION', 'Turn warm leads into clients', 'Stories, soft-CTA Reels', '1× per week'],
    ]
)

sub_heading(doc, 'Discovery Content — Best Performing Types')
bullet(doc, 'Before and after transformations (most shared format in interior design)')
bullet(doc, '"Unbelievable" reveals — rooms that look impossible to have changed')
bullet(doc, 'Renovation mistake reveals — "This is what happens with the wrong contractor"')
bullet(doc, 'Trend reactions — "Here\'s why I would never put this in a client\'s home"')
bullet(doc, '"Guess the budget" — show a finished room, ask followers to guess the cost')

sub_heading(doc, 'Trust Content — Best Performing Types')
bullet(doc, '"X mistakes people make when renovating their [room]"')
bullet(doc, '"How to choose the right [material/colour/contractor]"')
bullet(doc, 'Material mood boards with explanations of every decision')
bullet(doc, '"Project breakdown" carousels — one project, explained decision by decision')
bullet(doc, 'Behind-the-scenes process content — the thinking, not just the result')

sub_heading(doc, 'Conversion Content — Best Performing Types')
bullet(doc, 'Client testimonials (video preferred)')
bullet(doc, '"Here\'s exactly what working with us looks like" — full process walkthrough')
bullet(doc, '"We have X spots open for Q3" — real scarcity announcement')
bullet(doc, '"DM me a photo of your problem space and I\'ll tell you what I would do"')

section_heading(doc, '2.3  The Weekly Posting Schedule')
add_table(doc,
    ['Day', 'Format', 'Content Type', 'Purpose'],
    [
        ['Monday',    'Reel',           'Before/after transformation reveal',        'Maximum discovery reach'],
        ['Tuesday',   'Carousel 7–10s', '"5 mistakes / how to choose X"',            'Saves & trust building'],
        ['Wednesday', 'Reel',           'Process / material selection / BTS',        'Discovery + authenticity'],
        ['Thursday',  'Stories only',   'Polls, Q&A, behind-the-scenes',             'Relationship & loyalty'],
        ['Friday',    'Carousel or CTA','Trust carousel OR conversion post',         'Alternate weekly'],
        ['Saturday',  'Short Reel',     'Quick tip / design hack (7–15 seconds)',    'Shareability'],
        ['Sunday',    'Rest',           '—',                                         'Plan next week'],
    ]
)

section_heading(doc, '2.4  The Reel Formula')
body(doc, 'Every single Reel must follow this exact structure:')

sub_heading(doc, 'SECONDS 0–1.5 — THE HOOK  (most critical 1.5 seconds of your business)')
body(doc, '50% of viewers decide whether to continue within 1.5 seconds. Two components required:')
bullet(doc, 'Visual hook — movement from frame one, never start static')
bullet(doc, 'Text hook — 3 to 8 words creating irresistible need to know more')
body(doc, '')
body(doc, 'Hook formulas that work in Polish:')
bullet(doc, 'Curiosity:     "Ten błąd kosztuje ludzi tysiące złotych"')
bullet(doc, 'Curiosity:     "Czego żaden architekt ci nie powie"')
bullet(doc, 'Transformation:"To samo mieszkanie. 3 miesiące różnicy."')
bullet(doc, 'Pain point:    "Planujesz remont? Najpierw obejrzyj to."')
bullet(doc, 'Opinion:       "Szary jest przereklamowanym kolorem i mam dowody"')

sub_heading(doc, 'SECONDS 1.5–25 — THE CONTENT')
bullet(doc, 'Subtitles on everything — 85% of Instagram is watched without sound')
bullet(doc, 'Cut aggressively — every second without value removes viewers')
bullet(doc, 'Movement every 2–3 seconds — change shot, add text, use a transition')
bullet(doc, 'Face on camera when possible — +35% retention boost')

sub_heading(doc, 'LAST 3 SECONDS — THE CTA')
bullet(doc, '"Planujesz remont? Napisz do nas — link w bio"')
bullet(doc, '"Zapisz to na później" (for educational content — drives saves)')
bullet(doc, '"Wyślij to komuś, kto planuje mieszkanie" (drives DM shares)')

section_heading(doc, '2.5  The Carousel Formula')
add_table(doc,
    ['Slide', 'Content', 'Goal'],
    [
        ['Slide 1',       'Bold hook — claim or question that stops scrolling',         'Stop the scroll'],
        ['Slides 2–8',    'One specific point per slide. Short text + image.',           'Deliver value'],
        ['Slide 9',       'Tension — restate the stakes: why this matters',              'Keep engagement'],
        ['Slide 10',      'One CTA — save this / DM us / link in bio',                  'Convert'],
    ]
)

section_heading(doc, '2.6  Caption Strategy')
bullet(doc, 'Line 1: Hook — same energy as your Reel hook. This shows before "more."')
bullet(doc, 'Lines 2–5: The story — what was the problem? what decision did you make and why?')
bullet(doc, 'Lines 6–8: The teaching moment — what can someone take away even if they never hire you?')
bullet(doc, 'Last line: Specific engagement question or CTA')
body(doc, '')
callout(doc, '★ HASHTAG RULE:', 'Place hashtags in the FIRST COMMENT, not in the caption. Use 10–15 hashtags. Never use #interiordesign (200M+ posts — you will drown in noise).')

add_table(doc,
    ['Hashtag Tier', 'Post Volume', 'Examples', 'How Many'],
    [
        ['Niche',  'Under 100K',       '#projektowaniewnetrzwarszawa  #wnetrzaminimalistyczne', '4–5'],
        ['Mid',    '100K – 500K',      '#projektwnetrz  #dekoracjewnetrz  #architekturawnetrz', '4–5'],
        ['Broad',  '500K – 2M',        '#wnetrza  #interiordesignpoland  #domoweinspiracje',    '2–3'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 3
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 3 — GROWTH TACTICS')

section_heading(doc, '3.1  The Daily 30-Minute Growth Ritual')
body(doc, 'Do this every single day before you post anything. Skipping this is the single biggest growth mistake small accounts make.')

sub_heading(doc, '10 Minutes — Comment on Competitor Accounts')
body(doc, 'Go to your 5 target large Polish interior design accounts. Find their 3 most recent posts. Leave a genuine, specific, insightful comment. Not "Piękne ✨" — something like: "Wybór tkanin do tej sofy to strzał w dziesiątkę — ta faktura idealnie rozbija zimne tony ściany." People reading that comment will click your profile. Some will follow you.')

sub_heading(doc, '10 Minutes — Engage With Your Own Followers')
body(doc, 'Reply to every comment on your posts. Reply to every DM within the same day. Like and comment on stories from followers who engage with yours. Followers who feel seen become loyal followers who tell others.')

sub_heading(doc, '10 Minutes — Tag and Engage With Brands')
body(doc, 'Every material, piece of furniture, or paint colour in your projects — tag the brand. Comment on their latest post mentioning you used their product. Brands repost this constantly. One repost from a brand with 50,000 followers can bring 200 new followers to you in a single day.')

section_heading(doc, '3.2  The Collaboration Strategy')
body(doc, 'Collaboration is the fastest free growth strategy available to you.')

add_table(doc,
    ['Partner Type', 'Why Their Audience Matches Yours', 'Collaboration Format'],
    [
        ['Real estate agents',    'Showing apartments to people who will soon renovate',       'Joint before/after of a listed apartment'],
        ['Mortgage brokers',      'Buyers getting a mortgage will furnish and renovate',        'Cross-promote to same buying-stage audience'],
        ['Furniture/tile brands', 'Interior design fans who follow the brands you use',         'Tag → they reshare → you gain their audience'],
        ['Photographers',         'Design-conscious audience looking for portfolio inspiration', 'Credit exchange — they promote, you credit'],
        ['Contractors/builders',  'They need a designer to recommend to their clients',         'Referral relationship + cross-promotion'],
    ]
)

sub_heading(doc, 'DM Template for Proposing Collaborations')
callout(doc, 'COPY THIS:', '"Cześć [Name], jestem [Name] z OneDesign — projektujemy wnętrza w Warszawie. Zauważyłem/am że pracujemy dla podobnych klientów. Czy byłbyś/byłabyś zainteresowany/a wspólnym postem lub Stories? Mam konkretny pomysł jeśli chcesz pogadać."')
body(doc, 'Follow up once after 5 days. Never follow up twice.')

section_heading(doc, '3.3  Paid Promotion — Surgical Approach (500–1,500 PLN/month)')

add_table(doc,
    ['Rule', 'What to Do'],
    [
        ['Rule 1', 'Never boost a post that isn\'t already performing organically — wait 48 hours first'],
        ['Rule 2', 'Boost for Reach / Profile Visits — not Messages or Website Visits at this stage'],
        ['Rule 3', 'Target: Warsaw+30km + All Poland | Age: 26–45 | Interests: renovation, design, real estate'],
        ['Rule 4', '500 PLN: boost one top Reel per month for 10 days at 50 PLN/day'],
        ['Rule 5', 'Remaining 500–1,000 PLN: invest in ONE professional video shoot (= 4–8 Reels of content)'],
    ]
)

section_heading(doc, '3.4  Stories Strategy')
body(doc, 'Stories are chronically underused. They do three critical things: build the parasocial relationship, feed the algorithm with engagement data, and act as your most direct sales channel.')

add_table(doc,
    ['Day', 'Stories Content'],
    [
        ['Monday',    'Share Reel to Stories + poll ("Który styl wolisz — A czy B?")'],
        ['Wednesday', 'Behind the scenes from current project — raw, unedited, authentic'],
        ['Thursday',  'Full interactive sequence: question box + poll + personal/human moment'],
        ['Friday',    'Answer Thursday\'s questions — creates a loyal watch loop'],
        ['Saturday',  'Reshare Saturday\'s Reel'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 4
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 4 — THE MONETIZATION BLUEPRINT')

section_heading(doc, 'Four Revenue Streams — Built in This Order')
add_table(doc,
    ['Stream', 'When to Start', 'Realistic Monthly Income'],
    [
        ['1. Client Acquisition',  'Immediately',         '3–6 consultation leads → project revenue'],
        ['2. Affiliate Marketing', 'At 500 followers',    '200–500 PLN (growing to 3,000+ at 10K followers)'],
        ['3. Digital Products',    'At 1,000 followers',  '500–5,000 PLN depending on products created'],
        ['4. Brand Partnerships',  'At 5,000 followers',  '500–5,000 PLN per sponsored post'],
    ]
)

section_heading(doc, '4.1  Stream 1 — Client Acquisition')
bullet(doc, 'Every post ends with a soft CTA: "Planujesz remont? Napisz do nas — link w bio"')
bullet(doc, 'Weekly Story: "Wyślij mi zdjęcie swojego salonu przez DM — dam ci 3 konkretne sugestie"')
bullet(doc, 'After giving free advice in DMs: "Jeśli chcesz kompleksowe podejście, zapraszam na bezpłatną konsultację"')
body(doc, '')
callout(doc, '★ CONVERSION RATE:', 'If 20 people DM you from this tactic → 3–5 will book a consultation. That is a 15–25% conversion rate — far higher than any ad.')
body(doc, '')
body(doc, 'Set up Calendly (free) with a "15-minute free consultation" option. Put only this link in your bio. Remove everything else. The simpler the path, the more bookings you get.')

section_heading(doc, '4.2  Stream 2 — Affiliate Marketing')
body(doc, 'Every project you complete is a shopping opportunity for your audience. You earn 3–12% commission every time someone buys a product you recommended. The buyer pays the same price. You earn money for a recommendation you would have made anyway.')
body(doc, '')
bullet(doc, 'LTK (LikeToKnow.it): Primary platform for interior designers. Free to join. 8–15% commission rates.')
bullet(doc, 'Amazon.pl Storefront: Organise into collections — "My favourite lamps", "Best budget sofas", etc.')
bullet(doc, 'Direct affiliate deals: Contact your most-used Polish suppliers directly and ask if they offer commission.')
body(doc, '')
body(doc, 'Integration: Every project post should have "Wszystkie produkty z tego projektu znajdziecie w linku w bio."')

section_heading(doc, '4.3  Stream 3 — Digital Products')
body(doc, 'Build once. Sell forever.')

add_table(doc,
    ['Product', 'Description', 'Price', 'Platform'],
    [
        ['"Przewodnik Remontowy" PDF',   '20–30 page renovation guide. What to do, how to hire contractors, what mistakes to avoid.', '39–59 PLN',  'Gumroad'],
        ['Canva Moodboard Templates',    'Pack of 10 professional moodboard templates for homeowners to plan their spaces.',            '49–79 PLN',  'Etsy Poland / Gumroad'],
        ['"Jak urządzić bez architekta"','4–6 module video course: style, colour, layout, materials, lighting.',                       '199–349 PLN','Podia / Teachable'],
    ]
)

section_heading(doc, '4.4  Stream 4 — Brand Partnerships')
body(doc, 'At 5,000 followers with 3%+ engagement rate, start actively approaching brands.')
body(doc, '')
bullet(doc, 'Furniture: Black Red White, Agata Meble, Swarzędz Home')
bullet(doc, 'Paint: Śnieżka, Dulux, Farrow & Ball Poland, Beckers')
bullet(doc, 'Tile/Floor: Cerrad, Tubądzin, Paradyż')
bullet(doc, 'Kitchen: Nobilia, IKEA Partner Programme, Black Red White')
bullet(doc, 'Home accessories: Westwing, H&M Home Poland, Zara Home')
body(doc, '')
body(doc, 'When you reach 5,000 followers: create a 2-page media kit PDF with your stats, audience demographics, 3–4 best content examples, rates, and contact info. Send it proactively — brands will not come to you until you introduce yourself.')

# ══════════════════════════════════════════════════════════════════════════════
# PART 5
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 5 — THE 90-DAY EXECUTION PLAN')

section_heading(doc, 'Month 1 — Foundation (Weeks 1–4)')
sub_heading(doc, 'Weeks 1–2: Setup Only — No Growth Expectations Yet')
bullet(doc, 'Decide on your niche identity (Transformation or Smart Design Studio)')
bullet(doc, 'Rewrite your bio: who you help + what you do + differentiator + CTA')
bullet(doc, 'Switch to a Professional/Business Instagram account if not done')
bullet(doc, 'Create Linktree with: Calendly link + website portfolio link')
bullet(doc, 'Install CapCut — practice making one Reel from existing photos with voiceover')
bullet(doc, 'Build your hashtag bank: 20 niche/mid hashtags by room type and content type')
bullet(doc, 'Identify your 5 daily target competitor accounts for commenting')
bullet(doc, 'Identify 10 brands to tag regularly')
bullet(doc, 'Create Canva carousel template for consistent visual identity')
bullet(doc, 'Batch 8 posts (4 carousels + 4 Reel concepts from existing photos) — do not post yet')

sub_heading(doc, 'Weeks 3–4: Start Posting')
bullet(doc, 'Begin the daily 30-minute growth ritual every day')
bullet(doc, 'Post first Reel Monday — before/after format with voiceover from existing photos')
bullet(doc, 'Post first educational carousel Tuesday — "5 błędów przy urządzaniu salonu"')
bullet(doc, 'Maintain the 6-post weekly schedule for two full weeks')
bullet(doc, 'Reply to every comment within 24 hours')
bullet(doc, 'Start Thursday Stories interactive sequence')
bullet(doc, 'At Week 4: check Instagram Insights — which post got most reach, saves, shares? Make more of that.')
bullet(doc, 'Boost best Reel: 500 PLN / 10 days / Warsaw+Poland / age 26–45')

callout(doc, 'END OF MONTH 1 TARGET:', '+100–300 followers  ·  2–3 DM inquiries  ·  1–2 consultation requests')

section_heading(doc, 'Month 2 — Momentum (Weeks 5–8)')
bullet(doc, 'Book one professional video shoot at a completed project (500–800 PLN / 2 hours)')
bullet(doc, 'From shoot: produce 4 short Reels + 1 walkthrough + 15+ still photos (3–4 weeks of content)')
bullet(doc, 'Begin the weekly "DM your space" Story tactic')
bullet(doc, 'Launch carousel series: "Jak urządzić [salon/sypialnia/kuchnia] za [X] PLN" — one room per week')
bullet(doc, 'Reach out to 3 real estate agents or mortgage brokers with the collaboration DM template')
bullet(doc, 'Apply to LTK or set up Amazon.pl affiliate storefront')
bullet(doc, 'Begin linking products in every project post')
bullet(doc, 'Draft and publish "Przewodnik Remontowy" PDF on Gumroad (39–59 PLN)')
bullet(doc, 'Announce the digital product in Stories once per week')
bullet(doc, 'Boost one Reel at end of month: 500 PLN')

callout(doc, 'END OF MONTH 2 TARGET:', '+300–600 followers total  ·  4–6 DM inquiries/month  ·  First affiliate clicks  ·  First digital product sale')

section_heading(doc, 'Month 3 — Scale (Weeks 9–12)')
bullet(doc, 'Analyse your top 5 posts — what format, topic, hook performed best? Create 3 more versions of each.')
bullet(doc, 'Launch weekly recurring "Pytaj mnie o remont" — open question box Thursdays, answer Fridays')
bullet(doc, 'Execute first cross-account collab post')
bullet(doc, 'Create second digital product: Canva moodboard templates')
bullet(doc, 'Build email list: offer Przewodnik PDF free in exchange for email via Mailchimp')
bullet(doc, 'Create your 2-page media kit for future brand outreach')
bullet(doc, 'Full analytics review: follower growth rate, avg Reel reach, saves per carousel, bookings, affiliate earnings')
bullet(doc, 'Double down on every metric that is growing. Cut formats that are not.')

callout(doc, 'END OF MONTH 3 TARGET:', '800–2,500 followers  ·  4–8 consultation inquiries/month  ·  200–600 PLN/month affiliate income  ·  First digital product revenue')

# ══════════════════════════════════════════════════════════════════════════════
# PART 6
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 6 — MINDSET & THE VIRAL MOMENT')

section_heading(doc, '6.1  The Compounding Reality')
body(doc, 'Your first month will feel like shouting into a void. Your first Reel might get 200 views. This is normal. It is not a signal to stop. Instagram growth is exponential — but only after a critical mass of consistent content.')
body(doc, 'Every post you make is a permanent asset. An educational carousel posted in Month 1 will continue to get saves for 18 months. A Reel that seemed to underperform can get picked up by the algorithm weeks later. Your content library compounds over time. Do not stop because it feels slow.')

section_heading(doc, '6.2  What Actually Makes an Account Explode')
body(doc, 'There is one thing that, when it happens, will grow your account faster than all tactics combined: a viral Reel. For an account your size that means 50,000–500,000+ views. You can\'t fully control it — but you can dramatically increase the probability.')
body(doc, '')
body(doc, 'Content most likely to be DM-shared (the actual viral mechanism):')
bullet(doc, 'Solves a specific relatable problem: "Twoje mieszkanie wygląda małe bo robisz ten błąd"')
bullet(doc, 'Reveals a surprising truth: "Dlaczego meble z IKEA wyglądają drożej niż są"')
bullet(doc, 'Validates something people secretly believe: "Remont pod klucz jest wart każdej złotówki i tu jest dlaczego"')
bullet(doc, 'So useful someone saves AND sends it: checklists, decision frameworks, "before you buy" guides')
callout(doc, 'TEST FOR EVERY POST:', 'Ask yourself: "Is there a specific person someone would send this to?" If yes — you\'re on the right track.')

section_heading(doc, '6.3  What to Do the Moment a Reel Goes Viral')
bullet(doc, 'Post a follow-up Reel within 48 hours — same topic, different angle. "You asked about X — here\'s what I didn\'t have time to say."')
bullet(doc, 'Reply to every comment within 24 hours — the comment section is full of potential clients.')
bullet(doc, 'Pin the viral Reel to your profile — it\'s the first thing new visitors will see.')
bullet(doc, 'Update your bio — make the consultation booking path as clear and direct as possible.')
bullet(doc, 'Do NOT change your strategy. Go back to your schedule. Keep building.')

section_heading(doc, '6.4  Metrics That Matter vs. Metrics That Don\'t')
add_table(doc,
    ['Metric', 'Matters?', 'Why'],
    [
        ['Reel reach (non-followers)',    '★ YES',  'New people discovering you = growth'],
        ['Saves per post',                '★ YES',  'Quality of educational content signal'],
        ['Profile visits per week',       '★ YES',  'Interest generated by discovery content'],
        ['Consultation requests/month',   '★ YES',  'Direct business conversion'],
        ['Follower growth rate (weekly)', '★ YES',  'Momentum indicator'],
        ['Total follower count',          'No',     'Vanity metric at small scale'],
        ['Likes',                         'No',     'Weakest algorithmic signal'],
        ['Impressions',                   'No',     'Includes existing followers — not growth'],
        ['Emoji-only comments',           'No',     'No value signal'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 7 — CONTENT BANK
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'PART 7 — CONTENT BANK: 50 READY-TO-USE IDEAS')

section_heading(doc, 'Discovery Reels — Use for Monday & Wednesday')
ideas_discovery = [
    'Before/after reveal of your most dramatic transformation',
    '"Worst to best" — a room that looked terrible, completely reinvented',
    'Time-lapse of a room going from empty to fully styled',
    'Your reaction to a poorly designed space (screenshot submission from followers)',
    '"Guess the budget" — show a finished room, ask followers to guess the cost',
    '"Expensive vs. budget" — two similar looks, one 3× the price — can they tell?',
    'A common design trend you think is overrated and exactly why',
    '"Never do this in a small apartment" — one specific mistake shown visually',
    'The detail nobody notices but makes the whole room work',
    'A project from your portfolio that was the most challenging to design',
]
for i, idea in enumerate(ideas_discovery, 1):
    bullet(doc, f'{i:02d}.  {idea}')

section_heading(doc, 'Trust Carousels — Use for Tuesday & Alternate Fridays')
ideas_trust = [
    '"5 mistakes people make when choosing paint colours"',
    '"How to make any room look twice as big without knocking walls down"',
    '"The questions you MUST ask before hiring a contractor"',
    '"Why your apartment feels cold even with warm furniture"',
    '"The one thing worth splurging on in every room — and it\'s not what you think"',
    '"How to choose the right sofa for your space — a step-by-step guide"',
    '"Why IKEA furniture can look expensive or cheap — the difference"',
    '"The renovation timeline nobody tells you about"',
    '"How to pick a colour palette if you have no idea where to start"',
    '"What every renovation contract must include — or don\'t sign it"',
    '"Open plan vs. closed kitchen — which is right for your lifestyle?"',
    '"How to calculate how much furniture will fit before you buy anything"',
    '"The lighting mistake that makes every room look worse"',
    '"How to style shelves so they don\'t look cluttered"',
    '"Why you should never start a renovation without a design plan"',
]
for i, idea in enumerate(ideas_trust, 11):
    bullet(doc, f'{i:02d}.  {idea}')

section_heading(doc, 'Conversion Content — Use for Alternate Fridays')
ideas_conv = [
    'A client testimonial — their words, their before/after, their result',
    '"Here\'s exactly what happens when you work with us — from first call to finished room"',
    '"We have X spots open for Q3 — here\'s how to apply"',
    'A project post showing the client\'s reaction to the reveal',
    '"The most common question we get — and our honest answer"',
]
for i, idea in enumerate(ideas_conv, 26):
    bullet(doc, f'{i:02d}.  {idea}')

section_heading(doc, 'Stories — Use for Thursdays & Throughout the Week')
ideas_stories = [
    'Question box: "What\'s your biggest renovation fear?"',
    'Poll: "Which layout would you choose for this room? A or B?"',
    'Slider: "How overwhelmed do you feel when thinking about renovating?"',
    'Quiz: "What interior style are you? Answer 3 questions."',
    'Day-in-the-life during a project site visit',
    '"We\'re choosing final materials today — which one would you pick?"',
    'A behind-the-scenes mistake that happened on a project and how you solved it',
    '"DM me your biggest design problem — I\'ll answer the best ones in Stories"',
]
for i, idea in enumerate(ideas_stories, 31):
    bullet(doc, f'{i:02d}.  {idea}')

section_heading(doc, 'Short Reels (7–15 seconds) — Use for Saturdays')
ideas_short = [
    '"One change that transforms any bathroom — takes 2 hours"',
    '"The cheapest upgrade for any kitchen that looks expensive"',
    '"Stop buying these things for your apartment"',
    '"3 plants that work in any interior regardless of light"',
    '"The 60-30-10 colour rule explained in 15 seconds"',
    '"Why your mirrors are making your apartment look smaller"',
    '"The one piece of furniture worth buying new, not second-hand"',
    '"Warm vs. cool white — and why it actually matters"',
    '"How to test a paint colour before committing to the whole wall"',
    '"The fastest way to make a bedroom feel more luxurious"',
    '"This one lighting change costs 200 PLN and changes everything"',
    '"Why interior designers always choose odd numbers when styling"',
]
for i, idea in enumerate(ideas_short, 39):
    bullet(doc, f'{i:02d}.  {idea}')

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
part_heading(doc, 'QUICK REFERENCE CARD')

add_table(doc,
    ['Topic', 'Formula / Rule'],
    [
        ['Bio formula',           'Who you help  +  What you do  +  Differentiator  +  CTA with link'],
        ['Reel formula',          'Hook (0–1.5s)  →  Content (subtitles, cut hard)  →  CTA (1 action)'],
        ['Carousel formula',      'Hook slide  →  Value slides (1 point each)  →  Tension  →  CTA slide'],
        ['Hashtag system',        '4–5 niche (<100K)  +  4–5 mid (100K–500K)  +  2–3 broad (<2M)'],
        ['Daily ritual',          '10 min comment competitors  +  10 min engage followers  +  10 min tag brands'],
        ['Budget split',          '500 PLN boost top Reel  +  500–1,000 PLN professional video shoot'],
        ['Posting schedule',      'Mon Reel / Tue Carousel / Wed Reel / Thu Stories / Fri Carousel / Sat Short Reel'],
        ['Revenue order',         'Client leads  →  Affiliates (500 flw)  →  Digital products (1K)  →  Brand deals (5K)'],
        ['The viral test',        'Ask: "Is there a specific person someone would DM this to?"'],
        ['Check analytics',       'Every Sunday — which posts drove most new profile visits + saves this week'],
    ]
)

# ── Final closing callout ─────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
shade_paragraph(p, '1A1A2E')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(p, before=8, after=8)
r = p.add_run('  Execute this plan consistently for 90 days before evaluating results.  ')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = GOLD
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
shade_paragraph(p2, '1A1A2E')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(p2, before=0, after=12)
r2 = p2.add_run('The biggest risk is not a bad strategy — it is inconsistent execution of a good one.')
r2.italic = True
r2.font.size = Pt(10)
r2.font.color.rgb = WHITE
r2.font.name = 'Calibri'

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\filip\Desktop\claudecode\onedesignpl_masterplan.docx'
doc.save(out)
print(f'Saved: {out}')
